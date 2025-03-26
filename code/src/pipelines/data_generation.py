import random
from email.mime.multipart import MIMEMultipart
from pathlib import Path
import yaml
from datetime import datetime, timedelta
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from email.utils import formataddr, formatdate
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
import io
import re

class DataGenerator:
    # ------------------------------
    # Configurations
    # ------------------------------
    project_root = Path(__file__).parent.parent.parent
    BORROWERS = ["ABTB Mid-Atlantic LLC", "XYZ Capital LLC", "Global Investments Inc.","Motion Industries PVT Limited"]
    BANKS = ["Citi Bank NA", "Wells Fargo Bank", "Bank of America", "JP Morgan"]
    LOAN_TYPES = ["Term Loan A-5", "Revolving Credit Facility", "Bridge Loan"]
    SHARE_ADJUSTMENT_REASONS = ["Corporate Action", "Voluntary Adjustment", "Regulatory Requirement"]
    CURRENCY_CHARS = ["USD", "$", "$$"]

    def __init__(self):
        with open(f"{self.project_root}/src/config/paths.yaml") as f:
            self.paths = yaml.safe_load(f)
            
        self.banks = ["Citi Bank NA", "Wells Fargo", "Bank of America"]
        self.loan_types = ["Term Loan A-5", "Revolving Credit", "Bridge Loan"]
        
    # ------------------------------
    # Entity Generation Utilities
    # ------------------------------

    def random_date(self,start_date="2023-01-01", end_date="2025-12-31"):
        start = datetime.strptime(start_date, "%Y-%m-%d")
        end = datetime.strptime(end_date, "%Y-%m-%d")
        random_date = start + timedelta(days=random.randint(0, (end - start).days))
        return random_date.strftime("%d-%b-%Y")

    def generate_entities(self) -> dict:
        """Generate structured entities for emails/attachments"""
        return {
            "borrower": random.choice(self.BORROWERS),
            "amount": f"{random.choice(self.CURRENCY_CHARS)} {random.randint(1_000_000, 10_000_000):,}",
            "effective_date": self.random_date(),
            "loan_id": f"LOAN-{random.randint(1000,9999)}",
            "account_number": str(random.randint(10**9, 10**10 - 1)),
            "aba_number": str(random.randint(10**8, 10**9 - 1)),  # New field  
            "action_required": random.choice(["FUND_TRANSFER", "NONE"])  # New field 
        }
    
    # ------------------------------
    # Email Content Generators
    # ------------------------------
    def generate_email_body(self,request_type: str, include_entities_in_body: bool = True) -> tuple:
        """Generate email body with guaranteed entity inclusion"""
        entities = self.generate_entities()
        
        # Base template with mandatory entities
        body = f"""
        {random.choice(self.BANKS)}
        Loan Agency Services
        Date: {datetime.now().strftime("%d-%b-%Y")}
        
        Borrower: {entities['borrower']}
        Amount: {entities['amount']}
        Effective Date: {entities['effective_date']}
        Loan ID: {entities['loan_id']}
        Account #: {entities['account_number']}
        """
        
        # Add request-specific content
        if request_type == "loan_approval":
            body += f"""
            Dear Client,
            We are pleased to approve your {random.choice(self.LOAN_TYPES)}.
            """
        elif request_type == "share_adjustment":
            body += f"""
            Notification: Shares adjusted due to {random.choice(self.SHARE_ADJUSTMENT_REASONS)}.
            """
        
        # Add structured noise (preserves entity positions)
        if random.random() < 0.3:  # 30% chance of non-entity noise
            noise_options = [
                "\n[CONFIDENTIAL] This email contains privileged information.",
                "\nPlease contact your relationship manager for details.",
                "\nThis message is automatically generated."
            ]
            body += random.choice(noise_options)
        
        return body.strip(), entities

    # ------------------------------
    # Attachment Generators
    # ------------------------------
    def generate_pdf_attachment(self,entities: dict) -> bytes:
        """Generate PDF with structured entity table"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        
        # Create table
        data = [
            ["Field", "Value"],
            ["Borrower", entities['borrower']],
            ["Amount", entities['amount']],
            ["Effective Date", entities['effective_date']],
            ["Loan ID", entities['loan_id']]
        ]
        
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('ALIGN', (0,0), (-1,-1), 'CENTER')
        ]))
        
        doc.build([table])
        buffer.seek(0)
        return buffer.read()

    def generate_docx_attachment(self,entities: dict) -> bytes:
        """Generate DOCX with structured entities"""
        doc = Document()
        doc.add_heading('Loan Details', 0)
        
        doc.add_paragraph(f"Borrower: {entities['borrower']}")
        doc.add_paragraph(f"Amount: {entities['amount']}")
        doc.add_paragraph(f"Effective Date: {entities['effective_date']}")
        doc.add_paragraph(f"Account #: {entities['account_number']}")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()


    def generate_email(self) -> MIMEMultipart:
        """Generate a complete email with entities"""
        msg = MIMEMultipart()
        request_type = random.choice(["loan_approval", "share_adjustment"])
        
        # Generate content
        body, entities = self.generate_email_body(request_type)
        
        # Add headers
        msg['From'] = formataddr(("Banking Team", "operations@bank.com"))
        msg['To'] = "client@example.com"
        msg['Subject'] = f"{request_type.replace('_', ' ').title()} - {entities['loan_id']}"
        msg.attach(MIMEText(body))
        
        # Add attachments (80% chance)
        if random.random() < 0.8:
            # Add PDF
            pdf = MIMEApplication(
                self.generate_pdf_attachment(entities),
                _subtype="pdf"
            )
            pdf.add_header('Content-Disposition', 'attachment', filename="loan_details.pdf")
            msg.attach(pdf)
            
            # Add DOCX (50% chance)
            if random.choice([True, False]):
                docx = MIMEApplication(
                    self.generate_docx_attachment(entities),
                    _subtype="vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                docx.add_header('Content-Disposition', 'attachment', filename="terms.docx")
                msg.attach(docx)
        
        return msg
        
    def generate_dataset(self, num_samples=1000):
        output_dir = Path(self.paths['data']['raw'])
        output_dir.mkdir(parents=True, exist_ok=True)
        
        for i in range(num_samples):
            email = self.generate_email()
            self._save_email(email, output_dir/f"email_{i}.eml")
            
    def _save_email(self, email: MIMEMultipart, filename: str):
        """Save email as .eml file"""
        with open(filename, 'w') as f:
            f.write(email.as_string())